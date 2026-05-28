"""Converte docs/contrato-social-aresvida.md em contrato-social-aresvida.docx"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE = Path(__file__).parent / "contrato-social-aresvida.md"
DOCX_FILE = Path(__file__).parent / "contrato-social-aresvida.docx"


# ── helpers ──────────────────────────────────────────────────────────────────

def set_page_format(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)


def add_page_numbers(doc):
    """Adds page numbers to the footer (center)."""
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    ins = OxmlElement("w:instrText")
    ins.set(qn("xml:space"), "preserve")
    ins.text = " PAGE "
    run._r.append(ins)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para_spacing(para, before=0, after=4, line=None):
    from docx.shared import Pt as _Pt
    para.paragraph_format.space_before = _Pt(before)
    para.paragraph_format.space_after  = _Pt(after)
    if line:
        para.paragraph_format.line_spacing = line


def apply_inline_bold(para, text, base_size=12, base_font="Times New Roman",
                      align=None):
    """Splits text on **...** markers and adds runs with proper bold."""
    if align:
        para.alignment = align
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        set_font(run, base_font, base_size, bold=(i % 2 == 1))


def style_paragraph(para, font_name="Times New Roman", font_size=12,
                    bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=0, space_after=6, line_spacing=None,
                    color=None, indent_left=None):
    para.alignment = align
    para_spacing(para, space_before, space_after, line_spacing)
    for run in para.runs:
        set_font(run, font_name, font_size, bold, italic, color)
    if indent_left:
        para.paragraph_format.left_indent = Cm(indent_left)


# ── table builder ─────────────────────────────────────────────────────────────

def build_table(doc, rows):
    """rows = list of lists of strings (first row = header)."""
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_header = (r_idx == 0)
            clean = cell_text.strip().lstrip(":").rstrip(":")
            apply_inline_bold(para, clean, base_size=10)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                if is_header:
                    run.font.bold = True
            # header row: light grey background
            if is_header:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "D9D9D9")
                tcPr.append(shd)

    doc.add_paragraph()  # spacing after table


# ── markdown parser ───────────────────────────────────────────────────────────

def parse_and_build(doc, md_text):
    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── skip horizontal rules ──
        if re.match(r"^---+$", line.strip()):
            i += 1
            continue

        # ── blank line ──
        if not line.strip():
            i += 1
            continue

        # ── table ──
        if line.strip().startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                raw = lines[i].strip().strip("|")
                # skip separator row (| :--- | :---: |)
                if re.match(r"^[\s|:\-]+$", raw):
                    i += 1
                    continue
                cells = [c.strip() for c in raw.split("|")]
                table_rows.append(cells)
                i += 1
            if table_rows:
                build_table(doc, table_rows)
            continue

        # ── blockquote (aviso legal) ──
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent  = Cm(1.0)
            para.paragraph_format.right_indent = Cm(1.0)
            para_spacing(para, 2, 4)
            apply_inline_bold(para, text, base_size=10)
            for run in para.runs:
                run.font.name   = "Times New Roman"
                run.font.size   = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── H1 ──
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_spacing(para, 0, 4)
            run = para.add_run(text)
            set_font(run, "Times New Roman", 18, bold=True)
            i += 1
            continue

        # ── H2 ──
        if line.startswith("## ") and not line.startswith("### "):
            text = line[3:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_spacing(para, 10, 4)
            run = para.add_run(text)
            set_font(run, "Times New Roman", 14, bold=True)
            i += 1
            continue

        # ── H3 ──
        if line.startswith("### "):
            text = line[4:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_spacing(para, 8, 2)
            run = para.add_run(text)
            set_font(run, "Times New Roman", 12, bold=True)
            i += 1
            continue

        # ── unordered list ──
        if line.startswith("- "):
            text = line[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent  = Cm(1.0)
            para.paragraph_format.space_after  = Pt(2)
            apply_inline_bold(para, text, base_size=11)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            i += 1
            continue

        # ── lettered sub-items  a)  b)  etc. ──
        if re.match(r"^[a-z]\)", line.strip()):
            text = line.strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent       = Cm(1.5)
            para.paragraph_format.first_line_indent = Cm(-0.5)
            para_spacing(para, 0, 2)
            apply_inline_bold(para, text, base_size=11)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            i += 1
            continue

        # ── paragraph marker §, bold-lead lines ──
        # Normal body paragraph
        text = line.strip()
        if not text:
            i += 1
            continue

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para_spacing(para, 0, 4)

        # detect indent for §Nº lines
        if text.startswith("**§"):
            para.paragraph_format.left_indent = Cm(0.5)

        apply_inline_bold(para, text, base_size=11)
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

        i += 1


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    md_text = MD_FILE.read_text(encoding="utf-8")

    doc = Document()
    set_page_format(doc)
    add_page_numbers(doc)

    # default Normal style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    parse_and_build(doc, md_text)

    doc.save(DOCX_FILE)
    print(f"Salvo: {DOCX_FILE}")


if __name__ == "__main__":
    main()
